import speech_recognition as sr
import spacy
from docx import Document
from datetime import datetime
import sys
class ClinicalVoiceDocumentation:
    def __init__(self):
        # Initialize speech recognizer
        self.recognizer = sr.Recognizer()
        # Load English language model for spaCy
        print("Loading NLP model...")
        self.nlp = spacy.load("en_core_web_sm")
        self.doc = Document()

    def record_audio(self, duration=None):
        """Record audio from microphone"""
        with sr.Microphone() as source:
            print("Adjusting for ambient noise...")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)

            if duration:
                print(f"Recording for {duration} seconds...")
                audio = self.recognizer.listen(source, timeout=duration)
            else:
                print("Listening... (Press Ctrl+C to stop)")
                audio = self.recognizer.listen(source)

        return audio

    def transcribe_audio(self, audio):
        """Convert audio to text using Google Speech Recognition"""
        try:
            text = self.recognizer.recognize_google(audio)
            return text
        except sr.UnknownValueError:
            print("Could not understand the audio")
            return None
        except sr.RequestError as e:
            print(f"Could not request results; {e}")
            return None

    def process_clinical_text(self, text):
        """Process text using spaCy NLP"""
        if not text:
            return None

        doc = self.nlp(text)

        # Extract clinical entities
        clinical_info = {
            "symptoms": [],
            "medications": [],
            "procedures": [],
            "measurements": []
        }

        for ent in doc.ents:
            if ent.label_ in ["SYMPTOM", "PROBLEM"]:
                clinical_info["symptoms"].append(ent.text)
            elif ent.label_ in ["TREATMENT", "MEDICATION"]:
                clinical_info["medications"].append(ent.text)
            elif ent.label_ in ["PROCEDURE"]:
                clinical_info["procedures"].append(ent.text)
            elif ent.label_ in ["QUANTITY"]:
                clinical_info["measurements"].append(ent.text)

        return clinical_info

    def generate_document(self, text, clinical_info):
        """Generate clinical documentation in Word format"""
        # Add header
        self.doc.add_heading('Clinical Documentation', 0)
        self.doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Add transcribed text
        self.doc.add_heading('Transcribed Notes', level=1)
        self.doc.add_paragraph(text)

        # Add processed clinical information
        self.doc.add_heading('Clinical Information', level=1)

        if clinical_info["symptoms"]:
            self.doc.add_heading('Symptoms', level=2)
            self.doc.add_paragraph(", ".join(clinical_info["symptoms"]))

        if clinical_info["medications"]:
            self.doc.add_heading('Medications', level=2)
            self.doc.add_paragraph(", ".join(clinical_info["medications"]))

        if clinical_info["procedures"]:
            self.doc.add_heading('Procedures', level=2)
            self.doc.add_paragraph(", ".join(clinical_info["procedures"]))

        if clinical_info["measurements"]:
            self.doc.add_heading('Measurements', level=2)
            self.doc.add_paragraph(", ".join(clinical_info["measurements"]))

        # Save document
        filename = f'clinical_doc_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        self.doc.save(filename)
        return filename


def main():
    cvd = ClinicalVoiceDocumentation()

    try:
        while True:
            print("\nPress Enter to start recording (or 'q' to quit)")
            user_input = input()

            if user_input.lower() == 'q':
                break

            # Record audio
            audio = cvd.record_audio()

            # Transcribe audio to text
            print("Transcribing...")
            text = cvd.transcribe_audio(audio)

            if text:
                print("\nTranscribed text:")
                print(text)

                # Process text with NLP
                print("\nProcessing clinical information...")
                clinical_info = cvd.process_clinical_text(text)

                # Generate documentation
                filename = cvd.generate_document(text, clinical_info)
                print(f"\nClinical documentation saved as: {filename}")

    except KeyboardInterrupt:
        print("\nExiting...")
        sys.exit(0)


if __name__ == "__main__":
    main()
